"""
Word document generation tool.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.utils.config import settings


class DocumentGenerator:
    """
    Generates Microsoft Word (.docx) documents.
    """

    @staticmethod
    def generate_document(
        title: str,
        content: str,
        output_filename: str,
    ) -> str:
        """
        Generate a formatted Word document.

        Returns
        -------
        str
            Path to generated document.
        """

        output_dir = Path(settings.OUTPUT_DIR)
        output_dir.mkdir(parents=True, exist_ok=True)

        output_path = output_dir / output_filename

        document = Document()

        # -----------------------------------------------------
        # Title
        # -----------------------------------------------------

        heading = document.add_heading(title, level=1)
        heading.style.font.size = Pt(18)

        # -----------------------------------------------------
        # Body
        # -----------------------------------------------------

        for line in content.split("\n"):

            line = line.strip()

            if not line:
                continue

            document.add_paragraph(line)

        # -----------------------------------------------------
        # Save
        # -----------------------------------------------------

        document.save(output_path)

        return str(output_path)